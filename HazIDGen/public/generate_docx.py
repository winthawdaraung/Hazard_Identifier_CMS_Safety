import json
import sys
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def merge_cells_horizontally(table, row_idx, start_col, end_col):
    """Merge cells horizontally in a table row"""
    try:
        row = table.rows[row_idx]
        # Get the first cell to merge into
        merged_cell = row.cells[start_col]
        
        # Merge subsequent cells
        for col_idx in range(start_col + 1, end_col + 1):
            if col_idx < len(row.cells):
                cell_to_merge = row.cells[col_idx]
                merged_cell.merge(cell_to_merge)
        
        return merged_cell
    except Exception as e:
        print(f"Error merging cells: {e}")
        return None
def merge_cells_vertically(table, start_row, col_idx):
    """Merge cells vertically in a table column"""
    try:
        # Get the first cell to merge into
        merged_cell = table.cell(start_row, col_idx)
        # Merge subsequent cells
        for row_idx in range(start_row + 1, len(table.rows)):
            cell_to_merge = table.cell(row_idx, col_idx)
            merged_cell.merge(cell_to_merge)
        return merged_cell
    except Exception as e:
        print(f"Error merging cells: {e}")
        return None

def format_paragraph(paragraph, font_size=11, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Apply consistent formatting to a paragraph"""
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(font_size)
        run.font.bold = bold
    paragraph.alignment = alignment

def format_table_cell(cell, font_size=10, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Apply consistent formatting to a table cell"""
    for paragraph in cell.paragraphs:
        format_paragraph(paragraph, font_size, bold, alignment)

def add_formatted_heading(doc, text, level=1, font_size=14, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Add a consistently formatted heading"""
    heading = doc.add_heading(text, level)
    format_paragraph(heading, font_size, bold=True, alignment=alignment)
    return heading

def add_formatted_paragraph(doc, text, font_size=11, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Add a consistently formatted paragraph"""
    paragraph = doc.add_paragraph(text)
    format_paragraph(paragraph, font_size, bold, alignment)
    return paragraph

def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph"""
    try:
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # Add hyperlink styling
        c = OxmlElement('w:color')
        c.set(qn('w:val'), "0563C1")
        rPr.append(c)
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        
        return hyperlink
    except Exception as e:
        print(f"Warning: Could not create hyperlink for {url}: {e}")
        # Fallback: just add text
        paragraph.add_run(text)
        return None

def format_date(date_string):
    """Format date string to DD/MM/YYYY"""
    if not date_string:
        return "N/A"
    try:
        date_obj = datetime.fromisoformat(date_string.replace('Z', '+00:00'))
        return date_obj.strftime("%d/%m/%Y")
    except:
        return date_string

def create_hazard_table(doc, category_name, hazards_data):
    """Create a table for hazard category with Subject, Details, Recommendations columns"""
    try:
        # Add category heading with consistent formatting
        heading = add_formatted_heading(doc, category_name, level=2, font_size=12)
        
        # Create table with 3 columns: Subject, Details, Recommendations
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'  # Professional table style
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths for better proportion
        table.columns[0].width = Inches(2.2)   # Subject - slightly narrower
        table.columns[1].width = Inches(2.8)   # Details
        table.columns[2].width = Inches(3.5)   # Recommendations - wider for more content
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Subject'
        hdr_cells[1].text = 'Details'
        hdr_cells[2].text = 'Recommendations'
        
        # Format header row with consistent styling
        for cell in hdr_cells:
            format_table_cell(cell, font_size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            
            # Add professional header background (darker blue)
            shading_elm = parse_xml(r'<w:shd {} w:fill="4F81BD"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
        
        # Add data rows for selected hazards only
        # Convert hazards_data to a list to preserve all entries with same name
        hazards_list = []
        
        for hazard_name, hazard_data in hazards_data.items():
            selected_status = hazard_data.get('selected', False)
            # Handle both boolean and string values for selected
            if selected_status == True or selected_status == 'true' or selected_status == 'True':
                hazards_list.append({
                    'name': hazard_name,
                    'data': hazard_data
                })
        
        # Sort by name to group similar hazards together
        hazards_list.sort(key=lambda x: x['name'])
        
        for hazard_entry in hazards_list:
            hazard_name = hazard_entry['name']
            hazard_data = hazard_entry['data']
            
            row_cells = table.add_row().cells
            
            # Subject (hazard name) - use the name from the data if available, otherwise use the key
            display_name = hazard_data.get('name', hazard_name)
            subject_para = row_cells[0].paragraphs[0]
            subject_run = subject_para.runs[0] if subject_para.runs else subject_para.add_run()
            subject_run.text = display_name
            subject_run.bold = True
            
            # Details
            details_text = hazard_data.get('details', 'N/A')
            row_cells[1].text = details_text
            
            # Recommendations - use user input or default from Excel data
            recommendations_text = hazard_data.get('recommendations', '')
            if not recommendations_text or recommendations_text.strip() == '':
                # If no user input, use the default from Excel data
                recommendations_text = hazard_data.get('defaultRecommendations', 'Standard safety measures apply')
            elif recommendations_text == 'No specific safety measures available.':
                # If it's the default placeholder, use a more appropriate text
                recommendations_text = 'Standard safety measures apply'
            
            row_cells[2].text = recommendations_text
        
        # Format data rows with consistent styling
        for row in table.rows[1:]:  # Skip header row
            for i, cell in enumerate(row.cells):
                if i == 0:  # Subject column - center align and bold
                    format_table_cell(cell, font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                else:  # Details and Recommendations - left align
                    format_table_cell(cell, font_size=10, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        
        # Add spacing after table
        add_formatted_paragraph(doc, '', font_size=11)
    except Exception as e:
        print(f"Error creating hazard table for {category_name}: {e}")

def create_header_with_logo(doc, data=None):
    """Create the header with CERN CMS Safety logo and metadata table"""
    try:
        print("Creating header with logo...")
        # Get the correct path to assets directory
        script_dir = os.path.dirname(os.path.abspath(__file__))
        assets_dir = os.path.join(script_dir, '..', 'src', 'assets')
        print(f"Assets directory: {assets_dir}")
        
        # Create a table for the header with 5 columns: Logos and Metadata
        header_table = doc.add_table(rows=1, cols=5)
        header_table.style = 'Table Grid'
        
        # Set table alignment
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        header_table.columns[0].width = Inches(3.0)   # CERN Logo and CMS Logo
        header_table.columns[1].width = Inches(0.5)   # Reference
        header_table.columns[2].width = Inches(0.5)   # EDMS
        header_table.columns[3].width = Inches(0.5)   # Rev.
        header_table.columns[4].width = Inches(0.5)   # Validity

        # Logos cell (left side)
        logo_cell = header_table.rows[0].cells[0]
        
        # Clear the cell first and create a new paragraph
        logo_cell.text = ""
        logo_paragraph = logo_cell.paragraphs[0]
        
        # Add CERN logo
        cern_logo_path = os.path.join(assets_dir, 'CERN_logo.png')
        print(f"Looking for CERN logo at: {cern_logo_path}")
        if os.path.exists(cern_logo_path):
            print("CERN logo found, adding to document...")
            cern_run = logo_paragraph.add_run()
            cern_run.add_picture(cern_logo_path, width=Inches(0.5), height=Inches(0.5))
            # Add a space or line break between logos
            logo_paragraph.add_run(" ")  # or use "\n" for line break
        else:
            print(f"Warning: CERN logo not found at {cern_logo_path}")
        
        # Add CMS logo
        cms_logo_path = os.path.join(assets_dir, 'Logo CMS Safety.png')
        print(f"Looking for CMS logo at: {cms_logo_path}")
        if os.path.exists(cms_logo_path):
            print("CMS logo found, adding to document...")
            cms_run = logo_paragraph.add_run()
            cms_run.add_picture(cms_logo_path, width=Inches(0.5), height=Inches(0.5))
        else:
            print(f"Warning: CMS logo not found at {cms_logo_path}")
        
        # Set logo paragraph alignment
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Metadata cells (right side)
        # Add metadata information from user input or defaults
        reference_cell = header_table.rows[0].cells[1]
        reference_value = data.get('reference', 'CMS-SIT-0000000') if data else 'CMS-SIT-0000000'
        reference_cell.text = f'Reference:\n{reference_value}'
        
        edms_cell = header_table.rows[0].cells[2]
        edms_value = data.get('edms', 'CMS-EDMS-0000000') if data else 'CMS-EDMS-0000000'
        edms_cell.text = f'EDMS:\n{edms_value}'
        
        rev_cell = header_table.rows[0].cells[3]
        rev_cell.text = 'Rev.:\n0.1'
        
        validity_cell = header_table.rows[0].cells[4]
        validity_value = data.get('validity', '31/12/2023') if data else '31/12/2023'
        validity_cell.text = f'Validity:\n{validity_value}'
        
        # Format all cells with consistent styling
        for i, cell in enumerate(header_table.rows[0].cells):
            if i == 0:  # Logo cell
                format_table_cell(cell, font_size=8, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            else:  # Metadata cells
                format_table_cell(cell, font_size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Set row height for better appearance
        header_table.rows[0].height = Inches(0.6)

        # doc.add_paragraph()
        
        print("Header created successfully!")
            
    except Exception as e:
        print(f"Error creating header: {e}")
        import traceback
        traceback.print_exc()

def create_hazard_definitions_table(doc, data):
    """Create the hazard definitions table with Hazard, Check, Definition, Ref. columns"""
    try:
        # Create table with 4 columns: Hazard, Check, Definition, Ref.
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'  # Professional table style
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths to match the original document structure
        table.columns[0].width = Inches(1.2)   # Hazard - slightly wider
        table.columns[1].width = Inches(0.6)   # Check - slightly wider for better visibility
        table.columns[2].width = Inches(3.5)   # Definition - wider for more content
        table.columns[3].width = Inches(0.8)   # Ref. - slightly wider
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Hazard'
        hdr_cells[1].text = 'Check'
        hdr_cells[2].text = 'Definition'
        hdr_cells[3].text = 'Ref.'
        
        # Format header row with consistent styling
        for cell in hdr_cells:
            format_table_cell(cell, font_size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            
            # Add professional header background (darker blue)
            shading_elm = parse_xml(r'<w:shd {} w:fill="4F81BD"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            # Set text color to white for better contrast
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = None  # This will make it white on dark background
        
        # Get data from the provided data object
        hazard_definitions = data.get('hazardDefinitions', [])
        selected_hazards = data.get('selectedHazards', [])
        
        # If no hazard definitions provided, log error and return
        if not hazard_definitions:
            print("ERROR: No hazard definitions found in data. Please ensure Excel data is loaded properly.")
            # Add a single row indicating the error
            row_cells = table.add_row().cells
            row_cells[0].text = "No Data"
            row_cells[1].text = ""
            row_cells[2].text = "Hazard definitions not loaded from Excel file"
            row_cells[3].text = "N/A"
            return
        
        print(f"Processing {len(hazard_definitions)} hazard definitions")
        print(f"Selected hazards: {selected_hazards}")
        print(f"Hazard definitions data: {hazard_definitions}")
        
        # Add data rows from Excel data
        for hazard_def in hazard_definitions:
            row_cells = table.add_row().cells
            
            # Hazard name
            hazard_name = hazard_def.get('hazard', hazard_def.get('Hazard', ''))
            row_cells[0].text = str(hazard_name)
            
            # Check - mark if selected (☒ for selected, ☐ for unselected)
            is_selected = False
            if selected_hazards:
                # Convert selected_hazards list to lowercase for comparison
                selected_hazards_lower = [str(h).lower() for h in selected_hazards]
                is_selected = hazard_name.lower() in selected_hazards_lower
            
            row_cells[1].text = '☒' if is_selected else '☐'
            if is_selected:
                # Make the checkmark bold
                for paragraph in row_cells[1].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Definition
            definition = hazard_def.get('definition', hazard_def.get('Definition', ''))
            row_cells[2].text = str(definition)
            
            # Ref. (Reference) - create hyperlink if it's a URL
            ref_text = hazard_def.get('ref', hazard_def.get('Ref.', ''))
            ref_para = row_cells[3].paragraphs[0]
            ref_para.clear()  # Clear existing content
            
            if ref_text and str(ref_text).startswith('http'):
                add_hyperlink(ref_para, str(ref_text), 'Link HSE')
            else:
                ref_para.text = str(ref_text) if ref_text else 'Link HSE'
        
        # Note: "Other Hazards" is now handled through the Excel data above
        # No need to add it manually as it should come from the hazardDefinitions
        
        # Format all content cells with consistent styling
        for row in table.rows[1:]:  # Skip header row
            for i, cell in enumerate(row.cells):
                if i == 1:  # Check column - center align
                    format_table_cell(cell, font_size=10, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                elif i == 0:  # Hazard column - center align and bold
                    format_table_cell(cell, font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                else:  # Other columns - left align
                    format_table_cell(cell, font_size=10, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        
        # print(f"Hazard definitions table created with {len(table.rows)} rows")
        
        # Add spacing after table
        doc.add_paragraph()
        
    except Exception as e:
        print(f"Error creating hazard definitions table: {e}")
        import traceback
        traceback.print_exc()

def generate_hazard_document(data, output_path):
    """Generate the complete hazard identification document following CERN template"""
    try:
        print(f"Starting document generation...")
        print(f"Output path: {output_path}")
        print(f"Data keys: {list(data.keys())}")
        
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Page 1: Title Page
        # Add header with logo
        create_header_with_logo(doc, data)
        
        # Document title
        title = add_formatted_heading(doc, 'Safety Report', level=0, font_size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Subtitle
        subtitle = add_formatted_heading(doc, 'Hazard Identification Process in Areas', level=1, font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Location info
        location_text = f"Building {data.get('building', 'XXXX')}/{data.get('room', 'XX-xxx')} {data.get('location', 'Meyrin/Prevessin/Point 5')}"
        location_para = add_formatted_paragraph(doc, location_text, font_size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Signature table
        sig_table = doc.add_table(rows=2, cols=3)
        sig_table.style = 'Light Grid Accent 1'
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        sig_hdr = sig_table.rows[0].cells
        sig_hdr[0].text = 'Prepared by:'
        sig_hdr[1].text = 'Checked by:'
        sig_hdr[2].text = 'Approved by:'
        
        # Content row
        sig_content = sig_table.rows[1].cells
        creator_formatted = f"{data.get('creatorName', '')} ({data.get('creatorDepartment', '')})"
        sig_content[0].text = creator_formatted
        sig_content[1].text = ''
        sig_content[2].text = ''
        
        # Format signature table
        for row in sig_table.rows:
            for cell in row.cells:
                format_table_cell(cell, font_size=10, bold=False)
        
        # Format header row as bold
        for cell in sig_table.rows[0].cells:
            format_table_cell(cell, font_size=10, bold=True)
        
        add_formatted_paragraph(doc, 'Distribution to:', font_size=11, bold=True)
        add_formatted_paragraph(doc, 'CMS Safety, Activity Responsible, TSO.', font_size=11)
        
        # Page 2: History of changes
        doc.add_page_break()
        # Add header with logo
        create_header_with_logo(doc, data)
        add_formatted_heading(doc, 'History of changes', level=1, font_size=14)
        history_table = doc.add_table(rows=2, cols=3)
        history_table.style = 'Light Grid Accent 1'
        
        history_hdr = history_table.rows[0].cells
        history_hdr[0].text = 'Rev.'
        history_hdr[1].text = 'Date'
        history_hdr[2].text = 'Description of changes'
        
        history_content = history_table.rows[1].cells
        history_content[0].text = '0.1'
        history_content[1].text = datetime.now().strftime("%d/%m/%Y")
        history_content[2].text = 'Creation of the document'
        
        # Format history table
        for row in history_table.rows:
            for cell in row.cells:
                format_table_cell(cell, font_size=10, bold=False)
        
        # Format header row as bold
        for cell in history_table.rows[0].cells:
            format_table_cell(cell, font_size=10, bold=True)
        
        # Page 3: Contacts and Useful Links
        doc.add_page_break()
        # Add header with logo
        create_header_with_logo(doc, data)
        add_formatted_heading(doc, '1 CONTACTS AND USEFUL LINKS', level=1, font_size=14)
        
        # Get contact data from the input data, with fallback to hardcoded values
        contact_data = data.get('contactData', {})
        web_contacts = contact_data.get('webContacts', [
            {'title': 'CERN HSE', 'url': 'https://hse.cern/', 'description': 'Website'},
            {'title': 'Contacts CMS Safety', 'url': 'https://cmssafety.web.cern.ch/who-are-we', 'description': 'group of CMS Safety referents'},
            {'title': 'CMS RP', 'url': 'https://cmssafety.web.cern.ch/radiation-protection', 'description': 'CMS radiation protection information'},
            {'title': 'CMS Safety Training and Access Requirements', 'url': 'https://cmssafety.web.cern.ch/training-and-access-requirements', 'description': 'all mandatory and recommended training'},
            {'title': 'CERN Learning Hub', 'url': 'https://lms.cern.ch/', 'description': 'for the catalogue and registration to available training courses'},
            {'title': 'ADaMS', 'url': 'http://adams.web.cern.ch/adams/', 'description': 'for access requests'},
            {'title': 'IMPACT', 'url': 'https://impact.cern.ch/impact/secure/', 'description': 'tool for the declaration of an activity'},
            {'title': 'TREC', 'url': 'https://cmmsx.cern.ch/SSO/trec/', 'description': 'system for tracing potentially radioactive equipment'},
            {'title': 'EDH SIT', 'url': 'https://edh.cern.ch/Document/SupplyChain/SIT', 'description': 'for Storage and/or internal transport requests'}
        ])
        
        email_contacts = contact_data.get('emailContacts', [
            {'email': 'Cms-safety@cern.ch', 'description': 'group of CMS Safety (TC, LEXGLIMOS, DLEXGLIMOS)'},
            {'email': 'Cms-safety-team@cern.ch', 'description': 'group of CMS Safety Team (LEXGLIMOS Office)'},
            {'email': 'Cms-rso@cern.ch', 'description': 'group of CMS Radiation Safety Officers (RSO, DRSO)'}
        ])
        
        # Add web contacts
        for contact in web_contacts:
            para = doc.add_paragraph('• ')
            add_hyperlink(para, contact['url'], contact['title'])
            para.add_run(f': {contact["description"]}')
        
        # Add email contacts
        for contact in email_contacts:
            para = doc.add_paragraph('• ')
            add_hyperlink(para, f'mailto:{contact["email"]}', contact['email'])
            para.add_run(f': {contact["description"]}')
        
        # Page 4: Hazards definitions
        doc.add_page_break()
        # Add header with logo
        create_header_with_logo(doc, data)
        add_formatted_heading(doc, '2 HAZARDS DEFINITIONS', level=1, font_size=14)
        
        definition_text = ("According to ISO 45001 a hazard is defined as a source capable of causing injury and ill health. "
                          "Hazards can include sources with the potential to cause harm or hazardous situations, "
                          "or circumstances with the potential for exposure leading to injury and ill health.")
        
        def_para = add_formatted_paragraph(doc, definition_text, font_size=11)
        for run in def_para.runs:
            run.italic = True
        
        # Add ISO link
        iso_para = add_formatted_paragraph(doc, '', font_size=11)
        add_hyperlink(iso_para, 'https://www.iso.org/obp/ui/fr/#iso:std:iso:45001:ed-1:v1:en', 'ISO 45001')
        
        add_formatted_paragraph(doc, '1. In the CHECK column of the table below, please check the hazards identified for the activity.', font_size=11)
        add_formatted_paragraph(doc, '2. For each identified hazard, please refer to the identification sheet by simply clicking on the corresponding paragraph.', font_size=11)
        
        # Create hazard definitions table
        create_hazard_definitions_table(doc, data)
        
        # Page 5: Your Area - Activity Summary Information
        doc.add_page_break()
        # Add header with logo
        create_header_with_logo(doc, data)
        add_formatted_heading(doc, '3 YOUR AREA', level=1, font_size=14)
        add_formatted_heading(doc, '3.1 ACTIVITY SUMMARY INFORMATION', level=2, font_size=12)
        
        # Activity summary table
        summary_table = doc.add_table(rows=7, cols=5)
        summary_table.style = 'Light Grid Accent 1'
        
        # Fill activity summary
        rows_data = [
            ['Title', data.get('title', 'Enter the name of the specific activity'), '', '', ''],
            ['Personnel', 'Name of the activity responsible:', data.get('responsiblePerson', 'Enter the name of the person leading activity'), '', ''],
            ['', 'Estimated number of participants:', data.get('participantCount', 'Enter the number of people performing the activity'), '', ''],
            ['Dates', 'Start date of the activity:', format_date(data.get('startDate')), f'Estimated end date:', format_date(data.get("endDate"))],
            ['Location', 'Building number and specific zone:', f"{data.get('building', '')}/{data.get('location', '')}", '', ''],
            ['', 'Location details:', f"{data.get('building', '')}/{data.get('room', '')}", '', ''],
            ['Support', 'CERN specific support (Group):', data.get('cernSupport', 'Enter the name of the CERN group'), 'CMS specific support:', data.get('cmsSupport', 'Enter CMS team')]
        ]
        
        for i, row_data in enumerate(rows_data):
            cells = summary_table.rows[i].cells
            
            # Fill in the cell data first
            for j, cell_data in enumerate(row_data):
                if j < len(cells) and str(cell_data).strip():  # Only fill non-empty cells
                    cells[j].text = str(cell_data)
            
            # Merge empty cells with the previous non-empty cell
            last_non_empty = 0
            for j in range(1, len(row_data)):
                if j < len(cells):
                    if not str(row_data[j]).strip():  # Empty cell
                        # Find the last non-empty cell to merge with
                        for k in range(j-1, -1, -1):
                            if str(row_data[k]).strip():
                                last_non_empty = k
                                break
                        # Don't merge here, just note the position
                    else:
                        last_non_empty = j
            
            # Apply merging for consecutive empty cells
            j = 0
            while j < len(row_data):
                if j < len(cells) and str(row_data[j]).strip():  # Non-empty cell
                    # Count consecutive empty cells after this one
                    empty_count = 0
                    for k in range(j + 1, len(row_data)):
                        if not str(row_data[k]).strip():
                            empty_count += 1
                        else:
                            break
                    
                    # Merge if there are empty cells
                    if empty_count > 0 and j + empty_count < len(cells):
                        try:
                            merge_cells_horizontally(summary_table, i, j, j + empty_count)
                        except Exception as e:
                            print(f"Could not merge cells in row {i}: {e}")
                    
                    j += empty_count + 1
                else:
                    j += 1
        # Set column widths
        summary_table.columns[0].width = Inches(2.0)
        summary_table.columns[1].width = Inches(2.0)
        summary_table.columns[2].width = Inches(2.0)
        summary_table.columns[3].width = Inches(1.0)
        summary_table.columns[4].width = Inches(1.0)
        
        # Apply vertical merging for Personnel and Location sections
        try:
            # Merge Personnel cells (rows 1-2, column 0)
            personnel_cell = summary_table.cell(1, 0)  # "Personnel" cell
            empty_personnel_cell = summary_table.cell(2, 0)  # Empty cell below
            personnel_cell.merge(empty_personnel_cell)
            
            # Merge Location cells (rows 4-5, column 0) 
            location_cell = summary_table.cell(4, 0)  # "Location" cell
            empty_location_cell = summary_table.cell(5, 0)  # Empty cell below
            location_cell.merge(empty_location_cell)
            
            print("Vertical merging completed successfully")
        except Exception as e:
            print(f"Error with vertical merging: {e}")
        
        # Format all cells with consistent styling
        for row in summary_table.rows:
            for i, cell in enumerate(row.cells):
                if i == 0:  # First column headers - bold
                    format_table_cell(cell, font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
                else:  # Content cells
                    format_table_cell(cell, font_size=10, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        
        # Documents section
        add_formatted_heading(doc, 'Existing documents (EDMS, Indico, …)', level=3, font_size=11)
        
        doc_table = doc.add_table(rows=3, cols=2)
        doc_table.style = 'Light Grid Accent 1'
        
        doc_rows = [
            ['Safety file:', data.get('safetyDocuments', 'Risk assessments, certificates, training records, VICs, etc.')],
            ['Technical documents:', data.get('technicalDocuments', 'Technical documents for tooling/equipment used')],
            ['Other useful documents:', data.get('otherDocuments', 'Procedures, instructions, task sheets, etc.')]
        ]
        
        for i, (label, content) in enumerate(doc_rows):
            cells = doc_table.rows[i].cells
            cells[0].text = label
            cells[1].text = content
        # HSE Documents
        add_formatted_heading(doc, 'Link with HSE (including HSE-RP)', level=3, font_size=11)
        
        hse_table = doc.add_table(rows=2, cols=2)
        hse_table.style = 'Light Grid Accent 1'
        
        hse_data = [
            ['Support by HSE on an already existing subject of activity:', data.get('hseSupport', 'HSE-RP, HSE inspections, etc. …')],
            ['Reference documents (if any):', data.get('referenceDocuments', 'Additional supporting documentation from HSE: reports, derogation requests, advice, etc.')]
        ]
        for i, (label, content) in enumerate(hse_data):
            cells = hse_table.rows[i].cells
            cells[0].text = label
            cells[1].text = content
            
        # Page 6: Activity Description
        doc.add_page_break()
        # Add header with logo
        create_header_with_logo(doc, data)
        add_formatted_heading(doc, '3.2 DESCRIPTION OF THE ACTIVITY', level=2, font_size=12)
        
        description_text = data.get('activityDescription', 'Further details about the activity...')
        add_formatted_paragraph(doc, description_text, font_size=11)
        
        # HSE Guideline note
        guideline_para = add_formatted_paragraph(doc, 'For the Section below, please consider to have a look to this valuable HSE Guideline: ', font_size=11)
        add_hyperlink(guideline_para, 'https://edms.cern.ch/document/1114042', 'https://edms.cern.ch/document/1114042')
        
        # Page 7: Hazard identification section
        doc.add_page_break()
        # Add header with logo
        create_header_with_logo(doc, data)
        add_formatted_heading(doc, 'IDENTIFICATION OF THE HAZARDS FOR YOUR ACTIVITY', level=1, font_size=14)
        
        # Process hazard details
        hazard_details = data.get('hazardDetails', {})
        
        if hazard_details:
            for category_name, category_data in hazard_details.items():
                if category_data and isinstance(category_data, dict):
                    # Check if any hazards are selected in this category
                    selected_hazards = {}
                    for k, v in category_data.items():
                        selected_status = v.get('selected', False)
                        if selected_status == True or selected_status == 'true' or selected_status == 'True':
                            selected_hazards[k] = v
                    
                    if selected_hazards:
                        create_hazard_table(doc, category_name, selected_hazards)
        
        # The "Other Hazards" category is already processed above in the main loop
        # No additional processing needed since it's now part of hazardDetails
        
        if not hazard_details:
            doc.add_paragraph('No hazard details provided.')
        
        # Annex
        add_formatted_heading(doc, 'ANNEX: PICTURES', level=1, font_size=14)
        add_formatted_paragraph(doc, 'Attached to EDMS Reference.', font_size=11)
        
        # Save document
        print(f"Saving document to: {output_path}")
        doc.save(output_path)
        print(f"SUCCESS: Document generated successfully: {output_path}")
        return True
    except Exception as e:
        print(f"Error generating document: {e}")
        import traceback
        traceback.print_exc()
        raise e

def main():
    """Main function to handle command line arguments"""
    print(f"Python script started with args: {sys.argv}")
    print(f"Current working directory: {os.getcwd()}")
    
    if len(sys.argv) != 3:
        print("Usage: python generate_docx.py <input_json_file> <output_docx_file>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    print(f"Input file: {input_file}")
    print(f"Output file: {output_file}")
    
    try:
        # Check if input file exists
        if not os.path.exists(input_file):
            print(f"Error: Input file '{input_file}' not found.")
            sys.exit(1)
        
        # Check if we can write to output directory
        output_dir = os.path.dirname(output_file)
        if output_dir and not os.path.exists(output_dir):
            print(f"Error: Output directory '{output_dir}' does not exist.")
            sys.exit(1)
        
        with open(input_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        print(f"Successfully loaded JSON data with {len(data)} keys")
        generate_hazard_document(data, output_file)
        print("Document generation completed successfully!")
        
    except FileNotFoundError:
        print(f"Error: Input file '{input_file}' not found.")
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"Error: Invalid JSON in file '{input_file}': {e}")
        sys.exit(1)
    except Exception as e:
        print(f"Error generating document: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main() 